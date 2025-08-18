import uuid
from datetime import datetime
from typing import Any, Dict, List, Optional

from ..ports.llm import IChatLLM, IEmbedder
from ..ports.vector_store import IVectorStore
from .models import Document, QueryRequest, QueryResponse


class RagService:
    """Pure domain service orchestrating RAG operations."""

    def __init__(
        self, store: IVectorStore, llm: IChatLLM, embed: IEmbedder, rag_only: bool = False
    ):
        self.store = store
        self.llm = llm
        self.embed = embed
        self.rag_only = rag_only

    def query_documents(self, request: QueryRequest) -> QueryResponse:
        """Execute a RAG query with business logic isolated from infrastructure."""
        # Business rule: Generate query embedding
        query_vector = self.embed.embed_query(request.query)

        # Business rule: Search with tenant isolation
        hits = self.store.search(
            tenant_id=request.tenant_id, query_vector=query_vector, k=request.context_limit or 5
        )

        # Business rule: Generate answer using retrieved context
        answer = self.llm.answer(hits, request.query, rag_only=self.rag_only)

        # Business rule: Convert hits to domain documents
        source_documents = [
            Document(
                id=hit.get("id", ""),
                title=hit.get("source", "Unknown"),
                content=hit.get("text", ""),
                source=hit.get("source", ""),
                metadata={"score": hit.get("score", 0.0)},
                tenant_id=request.tenant_id,
                created_at=datetime.utcnow(),
            )
            for hit in hits
        ]

        return QueryResponse(
            answer=answer,
            sources=source_documents,
            confidence=self._calculate_confidence(hits),
            query_id=str(uuid.uuid4()),
            tenant_id=request.tenant_id,
        )

    def ingest_text(
        self, title: str, text: str, tenant_id: str, chunk_size: int = 800, overlap: int = 120
    ) -> Dict[str, Any]:
        """Ingest text with business logic for chunking and embedding."""
        # Business rule: Text chunking strategy
        chunks = self._chunk_text(text, chunk_size, overlap)

        # Business rule: Generate embeddings for all chunks
        embeddings = self.embed.embed_batch(chunks)

        # Business rule: Store with tenant isolation
        source_id = self.store.upsert_chunks(
            tenant_id=tenant_id, title=title, chunks=chunks, embeddings=embeddings
        )

        return {
            "success": True,
            "source_id": source_id,
            "chunks_created": len(chunks),
            "title": title,
            "tenant_id": tenant_id,
            "embedding_provider": self.embed.__class__.__name__,
        }

    def ingest_file(
        self, file_path: str, title: str, tenant_id: str, original_filename: str = None
    ) -> Dict[str, Any]:
        """Ingest a file with business logic for file processing."""
        import os

        # Business rule: Determine file type and processing strategy
        file_ext = os.path.splitext(file_path.lower())[1]

        try:
            # Business rule: Extract text based on file type
            if file_ext == ".txt" or file_ext == ".md":
                with open(file_path, "r", encoding="utf-8") as f:
                    text = f.read()
            elif file_ext in [".pdf", ".docx", ".doc"]:
                text = self._extract_document_text(file_path, file_ext)
            else:
                raise ValueError(f"Unsupported file type: {file_ext}")

            # Business rule: Use existing text ingestion logic
            result = self.ingest_text(title, text, tenant_id)
            result["file_type"] = file_ext
            result["original_filename"] = original_filename

            return result

        except Exception as e:
            return {
                "success": False,
                "error": str(e),
                "file_type": file_ext,
                "original_filename": original_filename,
            }

    def get_recent_documents(self, tenant_id: str, limit: int = 20) -> List[Dict[str, Any]]:
        """Get recently ingested documents for a tenant."""
        # Business rule: Retrieve recent documents with tenant isolation
        return self.store.get_recent_sources(tenant_id=tenant_id, limit=limit)

    def debug_query(self, query: str, tenant_id: str, k: int = 5) -> Dict[str, Any]:
        """Debug RAG query with detailed retrieval information."""
        # Business rule: Generate query embedding
        query_vector = self.embed.embed_query(query)

        # Business rule: Search with detailed scoring
        hits = self.store.search(tenant_id=tenant_id, query_vector=query_vector, k=k)

        # Business rule: Return debug information
        return {
            "query": query,
            "tenant_id": tenant_id,
            "embedding_model": self.embed.__class__.__name__,
            "retrieved_chunks": [
                {
                    "id": hit.get("id", ""),
                    "source": hit.get("source", "Unknown"),
                    "text": (
                        hit.get("text", "")[:200] + "..."
                        if len(hit.get("text", "")) > 200
                        else hit.get("text", "")
                    ),
                    "score": hit.get("score", 0.0),
                    "metadata": hit.get("metadata", {}),
                }
                for hit in hits
            ],
            "total_results": len(hits),
            "llm_model": self.llm.__class__.__name__,
            "rag_only_mode": self.rag_only,
        }

    def _extract_document_text(self, file_path: str, file_ext: str) -> str:
        """Business logic for extracting text from different document types."""
        try:
            if file_ext == ".pdf":
                # Try to import and use PyPDF2 or similar
                try:
                    import PyPDF2

                    with open(file_path, "rb") as file:
                        reader = PyPDF2.PdfReader(file)
                        text = ""
                        for page in reader.pages:
                            text += page.extract_text() + "\n"
                        return text
                except ImportError:
                    raise ValueError("PyPDF2 not installed. Cannot process PDF files.")

            elif file_ext in [".docx", ".doc"]:
                # Try to import and use python-docx
                try:
                    from docx import Document

                    doc = Document(file_path)
                    text = ""
                    for paragraph in doc.paragraphs:
                        text += paragraph.text + "\n"
                    return text
                except ImportError:
                    raise ValueError("python-docx not installed. Cannot process DOCX files.")

            else:
                raise ValueError(f"Unsupported document type: {file_ext}")

        except Exception as e:
            # Fallback: try to read as plain text
            try:
                with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                    return f.read()
            except Exception:
                raise ValueError(f"Failed to extract text from {file_ext} file: {str(e)}")

    def _chunk_text(self, text: str, chunk_size: int, overlap: int) -> List[str]:
        """Business logic for text chunking."""
        chunks = []
        i = 0
        while i < len(text):
            chunk_end = min(i + chunk_size, len(text))
            chunks.append(text[i:chunk_end])
            i += chunk_size - overlap
            if i >= len(text):
                break
        return chunks

    def _calculate_confidence(self, hits: List[Dict[str, Any]]) -> Optional[float]:
        """Business logic for confidence calculation."""
        if not hits:
            return 0.0

        # Simple confidence based on top hit score
        top_score = hits[0].get("score", 0.0)
        # Normalize score to 0-1 range (assuming similarity scores)
        return min(max(top_score, 0.0), 1.0)


class DocumentService:
    """Pure domain service for document management operations."""

    def __init__(self, store: IVectorStore):
        self.store = store

    def validate_document_access(
        self, document_id: str, tenant_id: str, user_tenant_id: str
    ) -> bool:
        """Business rule: Validate tenant access to documents."""
        return tenant_id == user_tenant_id

    def get_document_metadata(self, document_id: str, tenant_id: str) -> Optional[Dict[str, Any]]:
        """Get document metadata with business validation."""
        # Business logic would go here for retrieving document metadata
        # This would use additional ports for document storage
        return None


class TenantService:
    """Pure domain service for tenant management operations."""

    def validate_cross_tenant_access(
        self, user_role: str, user_tenant: str, target_tenant: str
    ) -> bool:
        """Business rule: Determine if user can access different tenant's data."""
        # Business rule: Only owners can cross tenant boundaries
        if user_role == "owner" and user_tenant != target_tenant:
            return True
        return user_tenant == target_tenant

    def get_tenant_limits(self, tenant_id: str) -> Dict[str, int]:
        """Business rule: Get tenant-specific limits."""
        # Default limits - could be customized per tenant
        return {"max_documents": 10000, "max_queries_per_hour": 1000, "max_chunk_size": 2000}
